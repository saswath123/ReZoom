import fitz  # PyMuPDF
from docx import Document
import os
import re

def extract_text_from_pdf(file_path):
    """Extract text from PDF with better error handling"""
    try:
        text = ""
        pdf = fitz.Document(file_path)
        
        for page_num in range(len(pdf)):
            page = pdf[page_num]
            page_text = page.get_text()  # type: ignore
            if page_text and page_text.strip():
                text += page_text + "\n"
        
        pdf.close()
        
        if not text or len(text.strip()) < 20:
            raise Exception("PDF contains insufficient text")
        
        # Clean the text
        text = re.sub(r'\n\s*\n', '\n\n', text)  # Remove excessive newlines
        text = re.sub(r'[^\x00-\x7F]+', ' ', text)  # Remove non-ASCII chars
        
        return text.strip()
        
    except Exception as e:
        raise Exception(f"PDF extraction failed: {str(e)}")

def extract_text_from_docx(file_path):
    """Extract text from DOCX with better error handling"""
    try:
        text = ""
        doc = Document(file_path)
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text and para.text.strip():
                text += para.text.strip() + "\n"
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text += " | ".join(row_text) + "\n"
        
        if not text or len(text.strip()) < 20:
            raise Exception("DOCX contains insufficient text")
        
        # Clean the text
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r'[^\x00-\x7F]+', ' ', text)
        
        return text.strip()
        
    except Exception as e:
        raise Exception(f"DOCX extraction failed: {str(e)}")

def extract_resume_text(file_path):
    """Main extraction handler with validation"""
    try:
        if not os.path.exists(file_path):
            raise Exception("File not found")
        
        if not os.path.getsize(file_path) > 0:
            raise Exception("File is empty")
        
        if file_path.lower().endswith(".pdf"):
            text = extract_text_from_pdf(file_path)
        elif file_path.lower().endswith(".docx"):
            text = extract_text_from_docx(file_path)
        else:
            raise Exception("Unsupported file format. Use PDF or DOCX")
        
        if not text or len(text) < 50:
            raise Exception("Could not extract sufficient text from resume")
        
        print(f"Successfully extracted {len(text)} characters of text")  # Debug
        return text
        
    except Exception as e:
        print(f"Extraction error: {str(e)}")  # Debug
        raise

def extract_profile_image_from_pdf(file_path):
    """Extract the first/largest image from PDF (likely the profile photo)"""
    try:
        pdf = fitz.Document(file_path)
        largest_image_bytes = None
        largest_area = 0
        largest_ext = "png"
        
        for page_num in range(len(pdf)):
            page = pdf[page_num]
            image_list = page.get_images(full=True)
            
            for img_info in image_list:
                xref = img_info[0]
                base_image = pdf.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]
                width = base_image.get("width", 0)
                height = base_image.get("height", 0)
                area = width * height
                
                # Filter out very small images
                if width >= 60 and height >= 60:
                    if area > largest_area:
                        largest_area = area
                        largest_image_bytes = image_bytes
                        largest_ext = image_ext
            
            # Stop if we found a suitable image on the first page
            if largest_image_bytes:
                break
                
        pdf.close()
        return largest_image_bytes, largest_ext
    except Exception as e:
        print(f"Failed to extract image from PDF: {e}")
        return None, None

def extract_profile_image_from_docx(file_path):
    """Extract the first/largest image from DOCX (likely the profile photo)"""
    try:
        from PIL import Image
        import io
        
        doc = Document(file_path)
        largest_image_bytes = None
        largest_size = 0
        
        for part in doc.part.package.iter_parts():
            if part.content_type.startswith("image/"):
                # Avoid thumbnail image
                if "thumbnail" in part.partname.lower():
                    continue
                image_bytes = part.blob
                size = len(image_bytes)
                if size > 5000:  # Filter out tiny icons
                    if size > largest_size:
                        largest_size = size
                        largest_image_bytes = image_bytes
                        
        if largest_image_bytes:
            img = Image.open(io.BytesIO(largest_image_bytes))
            ext = (img.format or "png").lower()
            return largest_image_bytes, ext
            
        return None, None
    except Exception as e:
        print(f"Failed to extract image from DOCX: {e}")
        return None, None

def extract_profile_image(file_path):
    """Extract profile image from resume (PDF/DOCX)"""
    if file_path.lower().endswith(".pdf"):
        return extract_profile_image_from_pdf(file_path)
    elif file_path.lower().endswith(".docx"):
        return extract_profile_image_from_docx(file_path)
    return None, None

